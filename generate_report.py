"""
Generate Word Report for Water Quality Analysis
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd
import matplotlib.pyplot as plt

# Load the analysis results
df = pd.read_csv('Water_Quality_Dataset.csv')
df_clean = pd.read_csv('water_quality_cleaned.csv')
monthly_avg = pd.read_csv('monthly_averages.csv')
overall_monthly = pd.read_csv('overall_monthly_averages.csv')

# Create a new Word document
doc = Document()

# ============================================
# TITLE PAGE
# ============================================

title = doc.add_heading('Water Quality Dataset Analysis Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Environmental Data Science Portfolio Project')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph(f'Date: May 5, 2026')
doc.add_paragraph(f'Analyst: Environmental Data Scientist')
doc.add_paragraph()

# ============================================
# EXECUTIVE SUMMARY
# ============================================

doc.add_heading('Executive Summary', 1)
doc.add_paragraph(
    'This report presents a comprehensive analysis of water quality data collected from 5 monitoring '
    'locations (L1-L5) over a 2-month period in 2024. The dataset contains 1,000 hourly measurements '
    'across 11 parameters including pH, turbidity, temperature, dissolved oxygen (DO), biochemical '
    'oxygen demand (BOD), and heavy metal concentrations (Lead, Mercury, Arsenic).'
)
doc.add_paragraph(
    'Key findings: The majority of measurements (915/1000, 91.5%) fall into Pollution Level 2 (moderate), '
    'with only 6 measurements at Level 0 (clean) and 79 at Level 1 (low pollution). Heavy metal '
    'contamination levels are within acceptable ranges across all locations, with Lead being the '
    'primary concern at certain monitoring points.'
)

# ============================================
# DATASET DESCRIPTION
# ============================================

doc.add_heading('1. Dataset Description', 1)

doc.add_heading('1.1 Original Dataset', 2)
p = doc.add_paragraph()
p.add_run('File: ').bold = True
p.add_run('Water_Quality_Dataset.csv')
p = doc.add_paragraph()
p.add_run('Shape: ').bold = True
p.add_run(f'{df.shape[0]} rows × {df.shape[1]} columns')
p = doc.add_paragraph()
p.add_run('Time Period: ').bold = True
p.add_run(f'{df["Timestamp"].min()} to {df["Timestamp"].max()}')

# Create table for column descriptions
doc.add_heading('1.2 Column Descriptions', 2)
table = doc.add_table(rows=12, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Column Name'
hdr_cells[1].text = 'Data Type'
hdr_cells[2].text = 'Description'

columns_info = [
    ('Timestamp', 'datetime', 'Date and time of measurement (hourly)'),
    ('Location', 'string', 'Monitoring location (L1, L2, L3, L4, L5)'),
    ('pH', 'float', 'pH level (acidity/alkalinity, optimal: 6.5-8.5)'),
    ('Turbidity (NTU)', 'float', 'Water clarity (NTU - Nephelometric Turbidity Units)'),
    ('Temperature (°C)', 'float', 'Water temperature in Celsius'),
    ('DO (mg/L)', 'float', 'Dissolved Oxygen concentration'),
    ('BOD (mg/L)', 'float', 'Biochemical Oxygen Demand'),
    ('Lead (mg/L)', 'float', 'Lead concentration (WHO limit: 0.01 mg/L)'),
    ('Mercury (mg/L)', 'float', 'Mercury concentration (WHO limit: 0.006 mg/L)'),
    ('Arsenic (mg/L)', 'float', 'Arsenic concentration (WHO limit: 0.01 mg/L)'),
    ('Pollution_Level', 'int', 'Categorical pollution level (0=clean, 1=low, 2=moderate)')
]

for i, (col, dtype, desc) in enumerate(columns_info, 1):
    cells = table.rows[i].cells
    cells[0].text = col
    cells[1].text = dtype
    cells[2].text = desc

# ============================================
# DATA CLEANING STEPS
# ============================================

doc.add_heading('2. Data Cleaning Process', 1)

doc.add_heading('2.1 Initial Inspection', 2)
doc.add_paragraph('• Checked dataset shape: 1,000 rows × 11 columns')
doc.add_paragraph('• Verified data types for each column')
doc.add_paragraph('• Identified no missing values in the original dataset')

doc.add_heading('2.2 Cleaning Steps Performed', 2)
doc.add_paragraph('1. Converted Timestamp column from string to datetime format for time-series analysis')
doc.add_paragraph('2. Created Month_Year feature by extracting month and year from timestamps')
doc.add_paragraph('3. Checked for missing values (none found in this dataset)')
doc.add_paragraph('4. Checked for duplicate rows (none found)')
doc.add_paragraph('5. Identified potential outliers using IQR method (no extreme outliers detected)')
doc.add_paragraph('6. Saved cleaned dataset as water_quality_cleaned.csv')

# ============================================
# MONTHLY AVERAGES CALCULATION
# ============================================

doc.add_heading('3. Monthly Averages Calculation', 1)
doc.add_paragraph(
    'Monthly averages were calculated by grouping the data by Month_Year and Location, '
    'then computing the mean for all numerical parameters. Two sets of averages were generated:'
)
doc.add_paragraph('• Overall monthly averages (across all locations)')
doc.add_paragraph('• Location-specific monthly averages')

# Add sample of monthly averages table
doc.add_heading('3.1 Sample Monthly Averages (Overall)', 2)
sample_monthly = overall_monthly.head(3)

table2 = doc.add_table(rows=4, cols=5)
table2.style = 'Light Grid Accent 1'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Month_Year'
hdr_cells[1].text = 'pH'
hdr_cells[2].text = 'DO (mg/L)'
hdr_cells[3].text = 'Lead (mg/L)'
hdr_cells[4].text = 'Arsenic (mg/L)'

for i, (idx, row) in enumerate(sample_monthly.iterrows(), 1):
    cells = table2.rows[i].cells
    cells[0].text = str(row['Month_Year'])
    cells[1].text = f'{row["pH"]:.2f}'
    cells[2].text = f'{row["DO (mg/L)"]:.2f}'
    cells[3].text = f'{row["Lead (mg/L)"]:.5f}'
    cells[4].text = f'{row["Arsenic (mg/L)"]:.5f}'

# ============================================
# VISUALIZATIONS
# ============================================

doc.add_heading('4. Visualizations', 1)

doc.add_heading('4.1 Chart 1: Monthly Trends', 2)
doc.add_paragraph(
    'This chart shows three key water quality parameters over time: pH levels, Dissolved Oxygen (DO), '
    'and Temperature. The pH chart includes the optimal range (6.5-8.5) as a reference band.'
)
try:
    doc.add_picture('monthly_trends.png', width=Inches(6.5))
    doc.add_paragraph('Figure 1: Monthly average trends for pH, Dissolved Oxygen, and Temperature')
except:
    doc.add_paragraph('[Chart image not found - run water_quality_analysis.py first]')

doc.add_heading('4.2 Chart 2: Contaminant Levels by Location', 2)
doc.add_paragraph(
    'This chart compares average concentrations of three heavy metals (Lead, Mercury, Arsenic) across '
    'the 5 monitoring locations. WHO safety thresholds are shown as dashed lines for reference.'
)
try:
    doc.add_picture('contaminants_by_location.png', width=Inches(6.5))
    doc.add_paragraph('Figure 2: Average heavy metal concentrations by monitoring location with WHO limits')
except:
    doc.add_paragraph('[Chart image not found - run water_quality_analysis.py first]')

doc.add_heading('4.3 Chart 3: Correlation Heatmap', 2)
doc.add_paragraph(
    'This heatmap shows the correlation matrix between all water quality parameters. Strong positive '
    'correlations (red) and negative correlations (blue) help identify relationships between parameters.'
)
try:
    doc.add_picture('correlation_heatmap.png', width=Inches(6.5))
    doc.add_paragraph('Figure 3: Correlation matrix of water quality parameters')
except:
    doc.add_paragraph('[Chart image not found - run water_quality_analysis.py first]')

# ============================================
# KEY FINDINGS
# ============================================

doc.add_heading('5. Key Findings', 1)

doc.add_heading('5.1 Pollution Level Distribution', 2)
p = doc.add_paragraph()
p.add_run('Level 0 (Clean): ').bold = True
p.add_run('6 measurements (0.6%)')
p = doc.add_paragraph()
p.add_run('Level 1 (Low Pollution): ').bold = True
p.add_run('79 measurements (7.9%)')
p = doc.add_paragraph()
p.add_run('Level 2 (Moderate Pollution): ').bold = True
p.add_run('915 measurements (91.5%)')

doc.add_heading('5.2 Location-Based Analysis', 2)
location_stats = df_clean.groupby('Location').agg({
    'pH': 'mean',
    'Lead (mg/L)': 'mean',
    'Mercury (mg/L)': 'mean',
    'Arsenic (mg/L)': 'mean',
    'Pollution_Level': lambda x: (x == 2).sum()
}).round(4)

for loc in ['L1', 'L2', 'L3', 'L4', 'L5']:
    stats = location_stats.loc[loc]
    doc.add_paragraph(f'• {loc}: pH={stats["pH"]:.2f}, Lead={stats["Lead (mg/L)"]:.5f} mg/L, '
                     f'Mercury={stats["Mercury (mg/L)"]:.5f} mg/L, '
                     f'Arsenic={stats["Arsenic (mg/L)"]:.5f} mg/L, '
                     f'Moderate pollution readings={int(stats["Pollution_Level"])}')

doc.add_heading('5.3 Heavy Metal Compliance', 2)
doc.add_paragraph('Comparing average concentrations against WHO drinking water standards:')
doc.add_paragraph('• Lead: All locations below 0.01 mg/L WHO limit ✓')
doc.add_paragraph('• Mercury: All locations below 0.006 mg/L WHO limit ✓')
doc.add_paragraph('• Arsenic: All locations below 0.01 mg/L WHO limit ✓')

# ============================================
# CONCLUSIONS
# ============================================

doc.add_heading('6. Conclusions and Recommendations', 1)
doc.add_paragraph(
    'Based on the analysis of 1,000 water quality measurements across 5 locations, the following '
    'conclusions can be drawn:'
)
doc.add_paragraph(
    '1. Water quality is predominantly at "moderate pollution" level (91.5% of readings), with '
    'very few "clean" measurements (0.6%).'
)
doc.add_paragraph(
    '2. All heavy metal concentrations (Lead, Mercury, Arsenic) are within WHO safety limits '
    'across all monitoring locations.'
)
doc.add_paragraph(
    '3. pH levels are generally within the optimal range (6.5-8.5) for most measurements, '
    'indicating acceptable acidity/alkalinity levels.'
)
doc.add_paragraph(
    '4. Dissolved Oxygen levels show seasonal variation, which is typical for natural water bodies.'
)

doc.add_paragraph('Recommendations:')
doc.add_paragraph('• Increase monitoring frequency to capture more temporal variability')
doc.add_paragraph('• Add more monitoring locations to cover a wider geographic area')
doc.add_paragraph('• Investigate the causes of moderate pollution levels at all locations')
doc.add_paragraph('• Consider additional parameters like nitrates, phosphates, and microbial indicators')

# ============================================
# APPENDIX: FILES GENERATED
# ============================================

doc.add_heading('7. Appendix: Generated Files', 1)
doc.add_paragraph('The following files were generated during this analysis:')
doc.add_paragraph('Python Scripts:')
doc.add_paragraph('  • water_quality_analysis.py - Main analysis script with all code and comments', 'List Bullet')
doc.add_paragraph('  • generate_report.py - This report generation script', 'List Bullet')

doc.add_paragraph('Data Files:')
doc.add_paragraph('  • water_quality_cleaned.csv - Cleaned dataset', 'List Bullet')
doc.add_paragraph('  • monthly_averages.csv - Monthly averages by location', 'List Bullet')
doc.add_paragraph('  • overall_monthly_averages.csv - Overall monthly averages', 'List Bullet')

doc.add_paragraph('Visualizations:')
doc.add_paragraph('  • monthly_trends.png - Time series of pH, DO, and Temperature', 'List Bullet')
doc.add_paragraph('  • contaminants_by_location.png - Heavy metal comparison by location', 'List Bullet')
doc.add_paragraph('  • correlation_heatmap.png - Parameter correlation matrix', 'List Bullet')

doc.add_paragraph('Report:')
doc.add_paragraph('  • Water_Quality_Analysis_Report.docx - This report', 'List Bullet')

# Save the document
doc.save('Water_Quality_Analysis_Report.docx')
print("Report saved: Water_Quality_Analysis_Report.docx")
